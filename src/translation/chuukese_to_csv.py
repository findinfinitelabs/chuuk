# chuukese_to_csv.py
import re
import csv
from docx import Document
from pathlib import Path

# Get the directory where this script is located
script_dir = Path(__file__).parent

# Load raw text from DOCX file
docx_path = script_dir / 'CHUUKESE_TO_ENGLISH.docx'
if not docx_path.exists():
    raise FileNotFoundError(f"DOCX file not found: {docx_path}")

print(f"ğŸ“– Reading from: {docx_path}")
doc = Document(docx_path)
raw = '\n'.join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])
print(f"ğŸ“„ Extracted {len(raw)} characters from DOCX")

# Find where the actual dictionary content starts
start_markers = ['ngang\t', 'ua\t', 'A\t', 'Ã¡Ã¡fengen']
start_para_index = -1
for i, para in enumerate(doc.paragraphs):
    para_text = para.text.strip()
    for marker in start_markers:
        if marker in para_text:
            start_para_index = i
            print(f"\nğŸ“ Found start marker '{marker}' in paragraph {i}")
            break
    if start_para_index != -1:
        break

if start_para_index == -1:
    print("\nâš ï¸ No clear start marker found in paragraphs, starting from beginning")
    start_para_index = 0

# Now parse the tab-delimited content properly
entries = []
entry_num = 1
pseudo_page = 2

# Process each paragraph that contains tabs as a potential dictionary entry, starting from start_para_index
for para in doc.paragraphs[start_para_index:]:
    text = para.text.strip()
    if not text or '\t' not in text:
        continue
    
    parts = text.split('\t')
    
    # Skip pronunciation guides, grammatical references, and obvious headers
    first_part = parts[0].strip()
    if (first_part.startswith('(') or 
        first_part.startswith('Stand') or
        first_part.startswith('Possessive') or
        first_part.startswith('Pronoun') or
        first_part.startswith('Example') or
        first_part.startswith('come go') or
        # Skip pronunciation guides (single letters or short entries with dashes)
        (' â€“ ' in first_part and len(first_part.split(' â€“ ')[0]) <= 5) or
        # Skip grammatical abbreviations
        first_part in ['n.', 'adj.', 'vt.', 'vi.', 'prep.', 'v.', 'int.'] or
        first_part.startswith('n. â€“') or
        first_part.startswith('adj. â€“') or
        first_part.startswith('vt. â€“') or
        first_part.startswith('vi. â€“') or
        first_part.startswith('prep. â€“') or
        # Skip entries without Chuukese characters (accents)
        not any(ord(c) > 127 for c in first_part) or
        len(first_part) < 2):
        continue
    
    # This looks like a dictionary entry
    entry = {
        'Entry #': entry_num,
        'Pseudo-Page': pseudo_page,
        'Chuukese Word / Form': first_part,
        'Part of Speech': '',
        'English Definition': '',
        'Examples': '',
        'Notes': ''
    }
    
    # Process the remaining parts
    definition_parts = []
    for part in parts[1:]:
        part = part.strip()
        if part:
            definition_parts.append(part)
    
    if definition_parts:
        entry['English Definition'] = '\t'.join(definition_parts)
    
    # Extract part of speech if present in the definition
    pos_patterns = ['v.', 'vt.', 'vi.', 'n.', 'adj.', 'adv.', 'int.', 'prep.']
    for pos in pos_patterns:
        if pos in entry['English Definition']:
            entry['Part of Speech'] = pos
            break
    
    entries.append(entry)
    entry_num += 1

print(f"\nğŸ“Š Processing summary:")
print(f"   Dictionary entries found: {len(entries)}")

if len(entries) > 0:
    print(f"\nğŸ“ First few entries:")
    for i, entry in enumerate(entries[:5]):
        chuuk = entry['Chuukese Word / Form']
        eng = entry['English Definition'][:100]
        print(f"   {i+1}. '{chuuk}' -> '{eng}{'...' if len(entry['English Definition']) > 100 else ''}'")
else:
    print("\nâš ï¸ No entries found")

# Ensure output directories exist
project_root = script_dir.parent.parent
training_data_dir = project_root / 'training_data'
uploads_dir = project_root / 'uploads'
training_data_dir.mkdir(exist_ok=True)
uploads_dir.mkdir(exist_ok=True)

# CSV file paths
training_csv_path = training_data_dir / 'chuukese_dictionary.csv'
uploads_csv_path = uploads_dir / 'chuukese_dictionary.csv'

fieldnames = ['Entry #', 'Pseudo-Page', 'Chuukese Word / Form', 'Part of Speech', 'English Definition', 'Examples', 'Notes']

# Save to training_data directory for model training (tab-delimited)
with open(training_csv_path, 'w', encoding='utf-8', newline='') as f:
    # Write header
    f.write('\t'.join(fieldnames) + '\n')
    # Write data rows
    for entry in entries:
        row = [
            str(entry['Entry #']),
            str(entry['Pseudo-Page']),
            entry['Chuukese Word / Form'],
            entry['Part of Speech'],
            entry['English Definition'],
            entry['Examples'],
            entry['Notes']
        ]
        f.write('\t'.join(row) + '\n')

# Save to uploads directory for database import via app (tab-delimited)
with open(uploads_csv_path, 'w', encoding='utf-8', newline='') as f:
    # Write header
    f.write('\t'.join(fieldnames) + '\n')
    # Write data rows
    for entry in entries:
        row = [
            str(entry['Entry #']),
            str(entry['Pseudo-Page']),
            entry['Chuukese Word / Form'],
            entry['Part of Speech'],
            entry['English Definition'],
            entry['Examples'],
            entry['Notes']
        ]
        f.write('\t'.join(row) + '\n')

print(f"âœ… Full table with {len(entries)} entries saved to:")
print(f"   ğŸ“š Training data: {training_csv_path}")
print(f"   ğŸ“¤ Uploads (for app): {uploads_csv_path}")
print(f"\nğŸ’¡ You can now:")
print(f"   1. Train models using the training_data CSV")
print(f"   2. Upload the CSV through the web app at http://localhost:5002")
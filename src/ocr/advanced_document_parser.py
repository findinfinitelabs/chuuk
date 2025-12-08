#!/usr/bin/env python3
"""
Advanced Document Parser for Complex Structured Documents
==========================================================

Handles large documents (200+ pages) with hierarchical structures,
indentation levels, formatting preservation, and intelligent content extraction.
"""

import os
import re
import json
import logging
from typing import Dict, List, Optional, Tuple, Any
from dataclasses import dataclass, asdict
from pathlib import Path
import xml.etree.ElementTree as ET

# Document processing imports
try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import fitz  # PyMuPDF
    PDF_ADVANCED_AVAILABLE = True
except ImportError:
    PDF_ADVANCED_AVAILABLE = False

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

@dataclass
class DocumentStructure:
    """Represents hierarchical document structure"""
    level: int
    text: str
    page_number: int
    paragraph_index: int
    formatting: Dict[str, Any]
    children: List['DocumentStructure'] = None
    parent_id: Optional[str] = None
    
    def __post_init__(self):
        if self.children is None:
            self.children = []

@dataclass
class ParsedDocument:
    """Complete parsed document with metadata"""
    total_pages: int
    total_paragraphs: int
    structure_tree: List[DocumentStructure]
    raw_text: str
    metadata: Dict[str, Any]
    processing_stats: Dict[str, Any]

class AdvancedDocumentParser:
    """
    Advanced parser for complex structured documents with preservation
    of hierarchical organization, formatting, and content relationships
    """
    
    def __init__(self, preserve_formatting: bool = True):
        self.preserve_formatting = preserve_formatting
        self.structure_patterns = {
            'heading_1': re.compile(r'^[A-Z][A-Z\s]+$'),  # ALL CAPS
            'heading_2': re.compile(r'^\d+\.\s+[A-Z]'),    # "1. Title"
            'heading_3': re.compile(r'^\d+\.\d+\s+'),      # "1.1 Subtitle"
            'bullet_point': re.compile(r'^[\-\•\*]\s+'),   # Bullet lists
            'numbered_list': re.compile(r'^\d+\)\s+'),     # "1) Item"
            'indented': re.compile(r'^\s{2,}'),           # Indented text
            'dictionary_entry': re.compile(r'^[^\s]+\s+[\-–]\s+'),  # "word - definition"
            'page_marker': re.compile(r'[Pp]age\s+\d+|^\d+$'),
        }
        
        self.formatting_indicators = {
            'bold': ['**', '__'],
            'italic': ['*', '_'],
            'underline': ['_'],
            'code': ['`', '```']
        }
        
    def parse_large_docx(self, file_path: str) -> ParsedDocument:
        """
        Parse large DOCX files with structure preservation
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            ParsedDocument with hierarchical structure
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for DOCX processing")
        
        logger.info(f"📖 Starting advanced parsing of: {os.path.basename(file_path)}")
        
        doc = Document(file_path)
        structure_tree = []
        current_level = 0
        parent_stack = []
        
        total_paragraphs = len(doc.paragraphs)
        processed_count = 0
        
        for para_idx, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            
            # Progress indicator
            if para_idx % 100 == 0:
                progress = (para_idx / total_paragraphs) * 100
                logger.info(f"📊 Processing: {progress:.1f}% ({para_idx}/{total_paragraphs} paragraphs)")
            
            if not text:
                continue
                
            # Extract formatting information
            formatting = self._extract_paragraph_formatting(paragraph)
            
            # Determine structure level and type
            level, structure_type = self._analyze_structure_level(text, formatting)
            
            # Skip page markers and headers/footers
            if self._is_page_marker(text):
                continue
            
            # Create structure node
            structure_node = DocumentStructure(
                level=level,
                text=text,
                page_number=self._estimate_page_number(para_idx, total_paragraphs),
                paragraph_index=para_idx,
                formatting=formatting
            )
            
            # Build hierarchy
            self._build_hierarchy(structure_node, parent_stack)
            structure_tree.append(structure_node)
            processed_count += 1
            
        # Generate processing statistics
        processing_stats = {
            'total_paragraphs_processed': processed_count,
            'structure_levels_found': len(set(node.level for node in structure_tree)),
            'avg_words_per_paragraph': sum(len(node.text.split()) for node in structure_tree) / max(processed_count, 1),
            'formatting_diversity': len(set(str(node.formatting) for node in structure_tree))
        }
        
        # Combine all text
        raw_text = '\n\n'.join(node.text for node in structure_tree)
        
        # Extract document metadata
        metadata = self._extract_document_metadata(doc, file_path)
        
        logger.info(f"✅ Parsing complete: {processed_count} paragraphs, {len(structure_tree)} structure nodes")
        
        return ParsedDocument(
            total_pages=self._estimate_total_pages(total_paragraphs),
            total_paragraphs=processed_count,
            structure_tree=structure_tree,
            raw_text=raw_text,
            metadata=metadata,
            processing_stats=processing_stats
        )
    
    def parse_large_pdf(self, file_path: str) -> ParsedDocument:
        """
        Parse large PDF files with structure preservation using PyMuPDF
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            ParsedDocument with hierarchical structure
        """
        if not PDF_ADVANCED_AVAILABLE:
            raise ImportError("PyMuPDF (fitz) is required for advanced PDF processing")
        
        logger.info(f"📄 Starting advanced PDF parsing of: {os.path.basename(file_path)}")
        
        doc = fitz.open(file_path)
        structure_tree = []
        para_idx = 0
        
        total_pages = len(doc)
        
        for page_num in range(total_pages):
            if page_num % 10 == 0:
                progress = (page_num / total_pages) * 100
                logger.info(f"📊 Processing page: {progress:.1f}% ({page_num + 1}/{total_pages})")
            
            page = doc.load_page(page_num)
            
            # Extract text with formatting information
            blocks = page.get_text("dict")
            
            for block in blocks["blocks"]:
                if "lines" not in block:
                    continue
                
                for line in block["lines"]:
                    line_text = ""
                    line_formatting = {}
                    
                    for span in line["spans"]:
                        text = span["text"].strip()
                        if text:
                            line_text += text + " "
                            # Extract formatting from span
                            if not line_formatting:  # Use first span's formatting
                                line_formatting = {
                                    'font': span.get('font', ''),
                                    'size': span.get('size', 0),
                                    'flags': span.get('flags', 0),
                                    'color': span.get('color', 0)
                                }
                    
                    line_text = line_text.strip()
                    if not line_text:
                        continue
                    
                    # Determine structure level
                    level, structure_type = self._analyze_structure_level(line_text, line_formatting)
                    
                    # Create structure node
                    structure_node = DocumentStructure(
                        level=level,
                        text=line_text,
                        page_number=page_num + 1,
                        paragraph_index=para_idx,
                        formatting=line_formatting
                    )
                    
                    structure_tree.append(structure_node)
                    para_idx += 1
        
        doc.close()
        
        # Generate processing statistics
        processing_stats = {
            'total_paragraphs_processed': para_idx,
            'pages_processed': total_pages,
            'avg_paragraphs_per_page': para_idx / max(total_pages, 1),
            'structure_levels_found': len(set(node.level for node in structure_tree))
        }
        
        # Combine all text
        raw_text = '\n\n'.join(node.text for node in structure_tree)
        
        # Extract metadata
        metadata = {
            'source_file': os.path.basename(file_path),
            'file_size_mb': os.path.getsize(file_path) / (1024 * 1024),
            'total_pages': total_pages,
            'parser_type': 'pymupdf_advanced'
        }
        
        logger.info(f"✅ PDF parsing complete: {total_pages} pages, {para_idx} paragraphs")
        
        return ParsedDocument(
            total_pages=total_pages,
            total_paragraphs=para_idx,
            structure_tree=structure_tree,
            raw_text=raw_text,
            metadata=metadata,
            processing_stats=processing_stats
        )
    
    def _extract_paragraph_formatting(self, paragraph) -> Dict[str, Any]:
        """Extract formatting information from DOCX paragraph"""
        formatting = {
            'alignment': str(paragraph.alignment) if paragraph.alignment else 'left',
            'style': paragraph.style.name if paragraph.style else 'Normal',
            'runs': []
        }
        
        for run in paragraph.runs:
            run_format = {
                'text': run.text,
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'font_name': run.font.name if run.font.name else 'default',
                'font_size': run.font.size.pt if run.font.size else 12
            }
            formatting['runs'].append(run_format)
        
        return formatting
    
    def _analyze_structure_level(self, text: str, formatting: Dict[str, Any]) -> Tuple[int, str]:
        """
        Analyze text and formatting to determine hierarchical level
        
        Returns:
            Tuple of (level, structure_type)
        """
        text_clean = text.strip()
        
        # Check for different structure patterns
        if self.structure_patterns['heading_1'].match(text_clean):
            return 1, 'heading_1'
        elif self.structure_patterns['heading_2'].match(text_clean):
            return 2, 'heading_2'
        elif self.structure_patterns['heading_3'].match(text_clean):
            return 3, 'heading_3'
        elif self.structure_patterns['bullet_point'].match(text_clean):
            return 4, 'bullet_point'
        elif self.structure_patterns['numbered_list'].match(text_clean):
            return 4, 'numbered_list'
        elif self.structure_patterns['dictionary_entry'].match(text_clean):
            return 5, 'dictionary_entry'
        elif self.structure_patterns['indented'].match(text):
            # Count indentation level
            indent_level = len(text) - len(text.lstrip())
            return min(5 + (indent_level // 4), 10), 'indented'
        
        # Check formatting-based levels (for PDF/advanced parsing)
        if isinstance(formatting, dict):
            font_size = formatting.get('size', 12)
            is_bold = formatting.get('flags', 0) & 2**4  # Bold flag in PyMuPDF
            
            if font_size > 16 or is_bold:
                return 1, 'heading_major'
            elif font_size > 14:
                return 2, 'heading_minor'
        
        return 6, 'body_text'
    
    def _is_page_marker(self, text: str) -> bool:
        """Check if text is a page marker or header/footer"""
        return bool(self.structure_patterns['page_marker'].match(text.strip()))
    
    def _estimate_page_number(self, para_idx: int, total_paragraphs: int) -> int:
        """Estimate page number based on paragraph position"""
        # Assume ~20 paragraphs per page on average
        return max(1, (para_idx // 20) + 1)
    
    def _estimate_total_pages(self, total_paragraphs: int) -> int:
        """Estimate total pages from paragraph count"""
        return max(1, total_paragraphs // 20)
    
    def _build_hierarchy(self, node: DocumentStructure, parent_stack: List[DocumentStructure]):
        """Build parent-child relationships in document structure"""
        # Remove parents that are at same or higher level
        while parent_stack and parent_stack[-1].level >= node.level:
            parent_stack.pop()
        
        # Set parent if available
        if parent_stack:
            node.parent_id = str(parent_stack[-1].paragraph_index)
            parent_stack[-1].children.append(node)
        
        # Add current node as potential parent
        parent_stack.append(node)
    
    def _extract_document_metadata(self, doc: Document, file_path: str) -> Dict[str, Any]:
        """Extract metadata from DOCX document"""
        try:
            core_props = doc.core_properties
            metadata = {
                'source_file': os.path.basename(file_path),
                'file_size_mb': os.path.getsize(file_path) / (1024 * 1024),
                'title': core_props.title or 'Unknown',
                'author': core_props.author or 'Unknown',
                'created': str(core_props.created) if core_props.created else 'Unknown',
                'modified': str(core_props.modified) if core_props.modified else 'Unknown',
                'parser_type': 'docx_advanced'
            }
        except Exception as e:
            metadata = {
                'source_file': os.path.basename(file_path),
                'file_size_mb': os.path.getsize(file_path) / (1024 * 1024),
                'parser_type': 'docx_advanced',
                'error': str(e)
            }
        
        return metadata
    
    def export_structure_json(self, parsed_doc: ParsedDocument, output_path: str):
        """Export parsed structure to JSON for analysis"""
        def serialize_structure(node: DocumentStructure) -> Dict:
            return {
                'level': node.level,
                'text': node.text[:200] + '...' if len(node.text) > 200 else node.text,
                'page_number': node.page_number,
                'paragraph_index': node.paragraph_index,
                'formatting': node.formatting,
                'children_count': len(node.children),
                'parent_id': node.parent_id
            }
        
        export_data = {
            'metadata': parsed_doc.metadata,
            'processing_stats': parsed_doc.processing_stats,
            'structure_tree': [serialize_structure(node) for node in parsed_doc.structure_tree]
        }
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(export_data, f, indent=2, ensure_ascii=False)
        
        logger.info(f"📄 Structure exported to: {output_path}")
    
    def create_training_dataset(self, parsed_doc: ParsedDocument) -> List[Dict[str, str]]:
        """
        Create training dataset from parsed document structure
        
        Returns:
            List of training examples for AI models
        """
        training_data = []
        
        for node in parsed_doc.structure_tree:
            # Extract potential dictionary entries
            if 'dictionary_entry' in str(node.level) or '–' in node.text or '—' in node.text:
                # Split on common separators
                separators = ['–', '—', ' - ', '  ', '\t']
                for sep in separators:
                    if sep in node.text:
                        parts = node.text.split(sep, 1)
                        if len(parts) == 2:
                            source = parts[0].strip()
                            target = parts[1].strip()
                            
                            if len(source) > 1 and len(target) > 2:
                                training_data.append({
                                    'source_text': source,
                                    'target_text': target,
                                    'context': f"Page {node.page_number}, Level {node.level}",
                                    'type': 'dictionary_entry'
                                })
                        break
            
            # Extract structured content for context learning
            if node.level <= 3 and len(node.text) > 10:
                training_data.append({
                    'source_text': f"Structure level {node.level}",
                    'target_text': node.text,
                    'context': f"Page {node.page_number}",
                    'type': 'structural_content'
                })
        
        logger.info(f"📚 Generated {len(training_data)} training examples")
        return training_data

# Usage example and CLI interface
def main():
    """CLI interface for advanced document parsing"""
    import argparse
    
    parser = argparse.ArgumentParser(description='Advanced Document Parser for Large Structured Documents')
    parser.add_argument('input_file', help='Path to input document (DOCX or PDF)')
    parser.add_argument('--output-dir', default='output', help='Output directory for results')
    parser.add_argument('--export-json', action='store_true', help='Export structure to JSON')
    parser.add_argument('--create-training-data', action='store_true', help='Create AI training dataset')
    
    args = parser.parse_args()
    
    # Create output directory
    output_dir = Path(args.output_dir)
    output_dir.mkdir(exist_ok=True)
    
    # Initialize parser
    doc_parser = AdvancedDocumentParser()
    
    # Parse document
    file_ext = Path(args.input_file).suffix.lower()
    if file_ext == '.docx':
        parsed_doc = doc_parser.parse_large_docx(args.input_file)
    elif file_ext == '.pdf':
        parsed_doc = doc_parser.parse_large_pdf(args.input_file)
    else:
        raise ValueError(f"Unsupported file type: {file_ext}")
    
    # Export results
    base_name = Path(args.input_file).stem
    
    if args.export_json:
        json_path = output_dir / f"{base_name}_structure.json"
        doc_parser.export_structure_json(parsed_doc, str(json_path))
    
    if args.create_training_data:
        training_data = doc_parser.create_training_dataset(parsed_doc)
        training_path = output_dir / f"{base_name}_training_data.json"
        with open(training_path, 'w', encoding='utf-8') as f:
            json.dump(training_data, f, indent=2, ensure_ascii=False)
        print(f"📚 Training data saved to: {training_path}")
    
    # Save raw text
    text_path = output_dir / f"{base_name}_extracted_text.txt"
    with open(text_path, 'w', encoding='utf-8') as f:
        f.write(parsed_doc.raw_text)
    
    print(f"✅ Processing complete! Results saved to: {output_dir}")
    print(f"📊 Statistics: {parsed_doc.processing_stats}")

if __name__ == "__main__":
    main()